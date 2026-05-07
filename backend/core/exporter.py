import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def export_to_text(content: str, filename: str = "transcription.txt") -> bytes:
    return content.encode('utf-8')


def export_to_pdf(content: str, filename: str = "transcription.pdf") -> bytes:
    try:
        from fpdf import FPDF

        class PDF(FPDF):
            def header(self):
                self.set_font('helvetica', 'B', 15)
                self.cell(0, 10, 'Voice2Knowledge Transcription', border=False, ln=True, align='C')
                self.ln(10)

            def footer(self):
                self.set_y(-15)
                self.set_font('helvetica', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', align='C')

        pdf = PDF()
        pdf.add_page()
        pdf.set_font('helvetica', '', 11)

        for line in content.split('\n'):
            if line.strip():
                pdf.multi_cell(0, 5, line)

        return pdf.output(dest='S').encode('latin-1')
    except ImportError:
        logger.warning("fpdf not installed, falling back to text export")
        return export_to_text(content, filename.replace('.pdf', '.txt'))
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        raise


def export_to_docx(content: str, filename: str = "transcription.docx") -> bytes:
    try:
        from docx import Document

        doc = Document()
        doc.add_heading('Voice2Knowledge Transcription', 0)

        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)

        stream = io.BytesIO()
        doc.save(stream)
        return stream.getvalue()
    except ImportError:
        logger.warning("python-docx not installed, falling back to text export")
        return export_to_text(content, filename.replace('.docx', '.txt'))
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        raise


def export_content(content: str, format: str, filename: str) -> bytes:
    format = format.lower()

    if format == 'txt' or format == 'text':
        return export_to_text(content, filename)
    elif format == 'pdf':
        return export_to_pdf(content, filename)
    elif format == 'docx':
        return export_to_docx(content, filename)
    else:
        raise ValueError(f"Unsupported export format: {format}")